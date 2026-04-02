"""RAG document store modeled after the nirma example store."""

from __future__ import annotations

import os
import re
from functools import lru_cache
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_ollama import OllamaEmbeddings

from ...models import DecisionDocumentChunk, RetrievedDocumentChunk

_TEXT_FILE_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".json", ".yaml", ".yml"}
_STOPWORDS = {
    "это",
    "как",
    "для",
    "при",
    "или",
    "что",
    "все",
    "если",
    "надо",
    "нужно",
    "быть",
    "также",
    "который",
    "которая",
    "которые",
    "через",
    "между",
    "должен",
    "должна",
    "должны",
    "будет",
    "план",
    "развития",
    "документ",
    "стратегия",
    "генеральный",
}


class DecisionDocumentStore:
    """Vector-backed in-memory store for one active planning document."""

    def __init__(self, *, mode: str = "elements") -> None:
        self._mode = mode
        self._embedding = get_document_embedding()
        self._store = InMemoryVectorStore(embedding=self._embedding)
        self._docs: list[Document] = []

    def add_document(
        self,
        *,
        path: str | None,
        document_name: str,
        document_text: str | None,
        document_type: str,
    ) -> list[DecisionDocumentChunk]:
        """Add one document to the vector store and return normalized chunks."""
        docs: list[Document]
        if path:
            docs = self._load_path_document(
                path=path,
                document_name=document_name,
                document_type=document_type,
            )
        else:
            docs = self._build_text_documents(
                text=str(document_text or "").strip(),
                document_name=document_name,
                document_type=document_type,
            )
        if not docs:
            raise ValueError("Не удалось сформировать чанки документа для RAG.")
        self._store.add_documents(docs)
        self._docs.extend(docs)
        return _chunks_from_docs(docs)

    def retrieve(self, *, query: str, top_k: int = 4) -> list[RetrievedDocumentChunk]:
        """Retrieve relevant chunks from the vector store."""
        normalized_query = str(query).strip()
        if not normalized_query:
            raise ValueError("RAG query не может быть пустым.")
        docs = self._retrieve_docs(normalized_query=normalized_query, top_k=int(top_k))
        score_lookup = self._retrieve_scores(
            normalized_query=normalized_query,
            top_k=max(int(top_k), 4),
        )
        retrieved: list[RetrievedDocumentChunk] = []
        for doc in docs:
            chunk_id = str(doc.metadata.get("chunk_id") or "").strip() or f"chunk-{len(retrieved)+1}"
            retrieved.append(
                RetrievedDocumentChunk(
                    chunk_id=chunk_id,
                    score=round(max(0.0, score_lookup.get(chunk_id, 0.0)), 4),
                    text=str(doc.page_content).strip(),
                    rationale=_retrieval_rationale(query=normalized_query, doc=doc),
                )
            )
        return retrieved

    def _retrieve_docs(self, *, normalized_query: str, top_k: int) -> list[Document]:
        """Retrieve documents with a robust fallback chain."""
        retrieval_errors: list[str] = []
        try:
            retriever = self._store.as_retriever(
                search_type="mmr",
                search_kwargs={"k": int(top_k)},
            )
            docs = retriever.invoke(normalized_query)
            if docs:
                return list(docs)
        except Exception as exc:
            retrieval_errors.append(_format_exception_text(exc))
        try:
            docs = self._store.max_marginal_relevance_search(normalized_query, k=int(top_k))
            if docs:
                return list(docs)
        except Exception as exc:
            retrieval_errors.append(_format_exception_text(exc))
        try:
            docs = self._store.similarity_search(normalized_query, k=int(top_k))
            if docs:
                return list(docs)
        except Exception as exc:
            retrieval_errors.append(_format_exception_text(exc))
        if retrieval_errors:
            details = "; ".join(item for item in retrieval_errors if item)
            raise RuntimeError(f"Не удалось выполнить RAG retrieval: {details or 'неизвестная ошибка'}")
        return []

    def _retrieve_scores(self, *, normalized_query: str, top_k: int) -> dict[str, float]:
        """Try to fetch scores, but never fail retrieval because of scoring."""
        try:
            scored_pairs = self._store.similarity_search_with_relevance_scores(
                normalized_query,
                k=int(top_k),
            )
        except Exception:
            return {}
        return {
            str(doc.metadata.get("chunk_id") or ""): float(score)
            for doc, score in scored_pairs
        }

    @property
    def doc_count(self) -> int:
        """Return the number of indexed chunks."""
        return len(self._docs)

    def _load_path_document(
        self,
        *,
        path: str,
        document_name: str,
        document_type: str,
    ) -> list[Document]:
        file_path = Path(str(path).strip()).expanduser()
        if not file_path.exists():
            raise ValueError(f"Файл документа не найден: {file_path}")
        if not file_path.is_file():
            raise ValueError(f"Путь не ведёт к файлу документа: {file_path}")
        suffix = file_path.suffix.lower()
        if suffix in _TEXT_FILE_SUFFIXES:
            text = file_path.read_text(encoding="utf-8").strip()
            return self._build_text_documents(
                text=text,
                document_name=document_name,
                document_type=document_type,
                source_ref=str(file_path),
            )

        if suffix == ".pdf":
            return self._load_pdf_document(
                file_path=file_path,
                document_name=document_name,
                document_type=document_type,
            )
        if suffix == ".docx":
            return self._load_docx_document(
                file_path=file_path,
                document_name=document_name,
                document_type=document_type,
            )
        raise ValueError(
            "Поддерживаются .txt/.md/.rst/.json/.yaml/.yml, .pdf и .docx."
        )

    def _load_pdf_document(
        self,
        *,
        file_path: Path,
        document_name: str,
        document_type: str,
    ) -> list[Document]:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError(
                "Для PDF RAG-документов установите pypdf."
            ) from exc

        reader = PdfReader(str(file_path))
        parts: list[str] = []
        for page in reader.pages:
            text = str(page.extract_text() or "").strip()
            if text:
                parts.append(text)
        text = "\n\n".join(parts).strip()
        if not text:
            raise ValueError("Не удалось извлечь текст из PDF-документа.")
        return self._build_text_documents(
            text=text,
            document_name=document_name,
            document_type=document_type,
            source_ref=str(file_path),
        )

    def _load_docx_document(
        self,
        *,
        file_path: Path,
        document_name: str,
        document_type: str,
    ) -> list[Document]:
        try:
            from docx import Document as WordDocument
        except ImportError as exc:
            raise RuntimeError(
                "Для DOCX RAG-документов установите python-docx."
            ) from exc

        word_document = WordDocument(str(file_path))
        parts: list[str] = []
        for paragraph in word_document.paragraphs:
            text = str(paragraph.text or "").strip()
            if text:
                parts.append(text)
        for table in word_document.tables:
            for row in table.rows:
                cells = [str(cell.text or "").strip() for cell in row.cells]
                row_text = " | ".join(value for value in cells if value)
                if row_text:
                    parts.append(row_text)
        text = "\n\n".join(parts).strip()
        if not text:
            raise ValueError("Не удалось извлечь текст из DOCX-документа.")
        return self._build_text_documents(
            text=text,
            document_name=document_name,
            document_type=document_type,
            source_ref=str(file_path),
        )

    def _build_text_documents(
        self,
        *,
        text: str,
        document_name: str,
        document_type: str,
        source_ref: str | None = None,
    ) -> list[Document]:
        raw_chunks = _build_document_chunks(text)
        docs: list[Document] = []
        for chunk in raw_chunks:
            docs.append(
                Document(
                    page_content=chunk.text,
                    metadata={
                        "chunk_id": chunk.chunk_id,
                        "document_name": document_name,
                        "document_type": document_type,
                        "source_ref": source_ref,
                        "char_start": chunk.char_start,
                        "char_end": chunk.char_end,
                        "keywords": list(chunk.keywords),
                    },
                )
            )
        return docs


@lru_cache(maxsize=1)
def get_document_embedding():
    """Build the shared Ollama embedding model for RAG documents."""
    _load_repo_env()
    model_name = os.environ.get("EMBEDDING_MODEL", "").strip() or "nomic-embed-text:latest"
    base_url = os.environ.get("BASE_URL", "").strip()
    if not base_url:
        raise RuntimeError("Set BASE_URL in .env before using RAG document embeddings.")
    return OllamaEmbeddings(
        model=model_name,
        base_url=base_url,
    )


def _load_repo_env() -> None:
    for parent in Path(__file__).resolve().parents:
        env_path = parent / ".env"
        if env_path.exists():
            load_dotenv(env_path, override=False)
            return


def _chunks_from_docs(docs: list[Document]) -> list[DecisionDocumentChunk]:
    chunks: list[DecisionDocumentChunk] = []
    for doc in docs:
        metadata = dict(doc.metadata or {})
        chunks.append(
            DecisionDocumentChunk(
                chunk_id=str(metadata.get("chunk_id") or f"chunk-{len(chunks)+1}"),
                text=str(doc.page_content).strip(),
                char_start=int(metadata.get("char_start") or 0),
                char_end=int(metadata.get("char_end") or len(str(doc.page_content))),
                keywords=[str(item) for item in metadata.get("keywords") or []],
            )
        )
    return chunks


def _normalize_text(text: str) -> str:
    return str(text).lower().replace("ё", "е")


def _tokenize(text: str) -> list[str]:
    words = re.findall(r"[a-zа-я0-9_]{3,}", _normalize_text(text))
    return [word for word in words if word not in _STOPWORDS]


def _top_keywords(text: str, *, limit: int = 12) -> list[str]:
    scores: dict[str, int] = {}
    for token in _tokenize(text):
        scores[token] = scores.get(token, 0) + 1
    ranked = sorted(scores.items(), key=lambda item: (-item[1], item[0]))
    return [token for token, _ in ranked[:limit]]


def _build_document_chunks(text: str, *, max_chars: int = 700, overlap_chars: int = 120) -> list[DecisionDocumentChunk]:
    paragraphs = [item.strip() for item in re.split(r"\n\s*\n+", str(text).strip()) if item.strip()]
    if not paragraphs:
        paragraphs = [str(text).strip()]
    chunks: list[DecisionDocumentChunk] = []
    cursor = 0
    buffer = ""
    chunk_start = 0
    for paragraph in paragraphs:
        candidate = f"{buffer}\n\n{paragraph}".strip() if buffer else paragraph
        if buffer and len(candidate) > max_chars:
            chunk_text = buffer.strip()
            chunks.append(
                DecisionDocumentChunk(
                    chunk_id=f"chunk-{len(chunks)+1}",
                    text=chunk_text,
                    char_start=chunk_start,
                    char_end=chunk_start + len(chunk_text),
                    keywords=_top_keywords(chunk_text),
                )
            )
            overlap_text = chunk_text[-overlap_chars:].strip()
            buffer = f"{overlap_text}\n\n{paragraph}".strip() if overlap_text else paragraph
            chunk_start = max(0, chunk_start + len(chunk_text) - len(overlap_text))
        else:
            if not buffer:
                chunk_start = cursor
            buffer = candidate
        cursor += len(paragraph) + 2
    if buffer.strip():
        chunk_text = buffer.strip()
        chunks.append(
            DecisionDocumentChunk(
                chunk_id=f"chunk-{len(chunks)+1}",
                text=chunk_text,
                char_start=chunk_start,
                char_end=chunk_start + len(chunk_text),
                keywords=_top_keywords(chunk_text),
            )
        )
    return chunks


def _retrieval_rationale(*, query: str, doc: Document) -> str:
    query_terms = set(_tokenize(query))
    chunk_terms = set(str(item) for item in (doc.metadata or {}).get("keywords") or [])
    overlap = sorted(query_terms & chunk_terms)
    if overlap:
        return "Совпали термины: " + ", ".join(overlap[:5])
    return "Возвращён MMR-ретривером как релевантный фрагмент."


def _format_exception_text(exc: Exception) -> str:
    text = str(exc).strip()
    if text:
        return text
    return exc.__class__.__name__
